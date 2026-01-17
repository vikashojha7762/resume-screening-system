"""
Resume Parser Service
Handles parsing of PDF, DOC/DOCX files and OCR for image-based resumes
"""
import logging
import re
from typing import Dict, List, Optional, Tuple, Any
from pathlib import Path
from io import BytesIO
import PyPDF2
import pdfplumber
from docx import Document
import pytesseract
from PIL import Image
import fitz  # PyMuPDF

logger = logging.getLogger(__name__)


class ResumeParser:
    """Parse resumes from various formats"""
    
    def __init__(self):
        self.section_patterns = {
            'experience': re.compile(r'(work\s+experience|employment|professional\s+experience|career)', re.IGNORECASE),
            'education': re.compile(r'(education|academic|qualifications|degrees)', re.IGNORECASE),
            'skills': re.compile(r'(skills|technical\s+skills|competencies|expertise)', re.IGNORECASE),
            'summary': re.compile(r'(summary|profile|objective|about)', re.IGNORECASE),
            'projects': re.compile(r'(projects|portfolio|work\s+samples)', re.IGNORECASE),
            'certifications': re.compile(r'(certifications|certificates|licenses)', re.IGNORECASE),
        }
    
    def parse(self, file_content: bytes, file_type: str, filename: str = "") -> Dict[str, Any]:
        """
        Parse resume file and extract structured data
        
        Args:
            file_content: Raw file bytes
            file_type: File extension (pdf, doc, docx, txt)
            filename: Original filename
            
        Returns:
            Dictionary with parsed resume data
        """
        try:
            logger.info(f"Parsing resume: {filename} (type: {file_type})")
            
            if file_type == 'pdf':
                return self._parse_pdf(file_content)
            elif file_type in ['doc', 'docx']:
                return self._parse_docx(file_content)
            elif file_type == 'txt':
                return self._parse_txt(file_content)
            else:
                raise ValueError(f"Unsupported file type: {file_type}")
                
        except Exception as e:
            logger.error(f"Error parsing resume {filename}: {str(e)}", exc_info=True)
            raise
    
    def _parse_pdf(self, file_content: bytes) -> Dict[str, Any]:
        """Parse PDF file using multiple methods for best results"""
        text_parts = []
        metadata = {}
        
        # Method 1: Try pdfplumber (better for tables and structured content)
        try:
            with pdfplumber.open(BytesIO(file_content)) as pdf:
                metadata['pages'] = len(pdf.pages)
                metadata['page_count'] = len(pdf.pages)
                
                for page_num, page in enumerate(pdf.pages, 1):
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
                    
                    # Extract tables
                    tables = page.extract_tables()
                    if tables:
                        metadata.setdefault('tables', []).extend(tables)
        except Exception as e:
            logger.warning(f"pdfplumber extraction failed: {str(e)}")
        
        # Method 2: Fallback to PyPDF2
        if not text_parts:
            try:
                pdf_file = BytesIO(file_content)
                pdf_reader = PyPDF2.PdfReader(pdf_file)
                metadata['pages'] = len(pdf_reader.pages)
                
                for page in pdf_reader.pages:
                    text = page.extract_text()
                    if text:
                        text_parts.append(text)
            except Exception as e:
                logger.warning(f"PyPDF2 extraction failed: {str(e)}")
        
        # Method 3: Try PyMuPDF as last resort
        if not text_parts:
            try:
                doc = fitz.open(stream=file_content, filetype="pdf")
                metadata['pages'] = len(doc)
                
                for page in doc:
                    text = page.get_text()
                    if text:
                        text_parts.append(text)
                doc.close()
            except Exception as e:
                logger.warning(f"PyMuPDF extraction failed: {str(e)}")
        
        # Method 4: OCR if text extraction failed
        if not text_parts:
            logger.info("Attempting OCR for PDF")
            return self._parse_with_ocr(file_content)
        
        full_text = "\n\n".join(text_parts)
        cleaned_text = self._clean_text(full_text)
        
        return {
            'raw_text': cleaned_text,
            'metadata': metadata,
            'sections': self._detect_sections(cleaned_text),
            'word_count': len(cleaned_text.split()),
            'char_count': len(cleaned_text)
        }
    
    def _parse_docx(self, file_content: bytes) -> Dict[str, Any]:
        """Parse DOCX file"""
        try:
            doc = Document(BytesIO(file_content))
            text_parts = []
            metadata = {}
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Extract tables
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                metadata.setdefault('tables', []).append(table_data)
            
            full_text = "\n".join(text_parts)
            cleaned_text = self._clean_text(full_text)
            
            return {
                'raw_text': cleaned_text,
                'metadata': metadata,
                'sections': self._detect_sections(cleaned_text),
                'word_count': len(cleaned_text.split()),
                'char_count': len(cleaned_text)
            }
        except Exception as e:
            logger.error(f"Error parsing DOCX: {str(e)}")
            raise
    
    def _parse_txt(self, file_content: bytes) -> Dict[str, Any]:
        """Parse plain text file"""
        try:
            text = file_content.decode('utf-8', errors='ignore')
            cleaned_text = self._clean_text(text)
            
            return {
                'raw_text': cleaned_text,
                'metadata': {},
                'sections': self._detect_sections(cleaned_text),
                'word_count': len(cleaned_text.split()),
                'char_count': len(cleaned_text)
            }
        except Exception as e:
            logger.error(f"Error parsing TXT: {str(e)}")
            raise
    
    def _parse_with_ocr(self, file_content: bytes) -> Dict[str, Any]:
        """Parse image-based resume using OCR"""
        try:
            # Convert PDF to images if needed
            doc = fitz.open(stream=file_content, filetype="pdf")
            text_parts = []
            
            for page_num in range(len(doc)):
                page = doc[page_num]
                pix = page.get_pixmap()
                img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                
                # Perform OCR
                ocr_text = pytesseract.image_to_string(img)
                if ocr_text:
                    text_parts.append(ocr_text)
            
            doc.close()
            
            full_text = "\n\n".join(text_parts)
            cleaned_text = self._clean_text(full_text)
            
            return {
                'raw_text': cleaned_text,
                'metadata': {'ocr': True, 'pages': len(text_parts)},
                'sections': self._detect_sections(cleaned_text),
                'word_count': len(cleaned_text.split()),
                'char_count': len(cleaned_text)
            }
        except Exception as e:
            logger.error(f"OCR parsing failed: {str(e)}")
            raise
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize extracted text"""
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove special characters but keep punctuation
        text = re.sub(r'[^\w\s\.\,\;\:\!\?\-\(\)]', ' ', text)
        
        # Remove multiple newlines
        text = re.sub(r'\n{3,}', '\n\n', text)
        
        # Strip leading/trailing whitespace
        text = text.strip()
        
        return text
    
    def _detect_sections(self, text: str) -> Dict[str, List[int]]:
        """
        Detect resume sections and return line numbers
        
        Returns:
            Dictionary mapping section names to line number ranges
        """
        lines = text.split('\n')
        sections = {}
        
        for line_num, line in enumerate(lines):
            for section_name, pattern in self.section_patterns.items():
                if pattern.search(line):
                    if section_name not in sections:
                        sections[section_name] = []
                    sections[section_name].append(line_num)
        
        return sections
    
    def extract_contact_info(self, text: str) -> Dict[str, Optional[str]]:
        """Extract contact information from resume text"""
        contact = {
            'email': None,
            'phone': None,
            'linkedin': None,
            'github': None,
            'website': None
        }
        
        # Email
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        emails = re.findall(email_pattern, text)
        if emails:
            contact['email'] = emails[0]
        
        # Phone
        phone_patterns = [
            r'[\+]?[(]?[0-9]{3}[)]?[-\s\.]?[0-9]{3}[-\s\.]?[0-9]{4,6}',
            r'\+?[1-9]\d{1,14}',
            r'\(\d{3}\)\s?\d{3}-\d{4}'
        ]
        for pattern in phone_patterns:
            phones = re.findall(pattern, text)
            if phones:
                contact['phone'] = phones[0]
                break
        
        # LinkedIn
        linkedin_pattern = r'(?:linkedin\.com/in/|linkedin\.com/pub/)([a-zA-Z0-9-]+)'
        linkedin = re.search(linkedin_pattern, text, re.IGNORECASE)
        if linkedin:
            contact['linkedin'] = f"linkedin.com/in/{linkedin.group(1)}"
        
        # GitHub
        github_pattern = r'(?:github\.com/)([a-zA-Z0-9-]+)'
        github = re.search(github_pattern, text, re.IGNORECASE)
        if github:
            contact['github'] = f"github.com/{github.group(1)}"
        
        # Website
        website_pattern = r'(?:https?://)?(?:www\.)?([a-zA-Z0-9-]+\.[a-zA-Z]{2,})'
        websites = re.findall(website_pattern, text)
        if websites:
            # Filter out common email domains
            excluded = ['gmail.com', 'yahoo.com', 'outlook.com', 'hotmail.com']
            for site in websites:
                if site.lower() not in excluded:
                    contact['website'] = site
                    break
        
        return contact


# Singleton instance
resume_parser = ResumeParser()

